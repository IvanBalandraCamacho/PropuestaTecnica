"""
Servicio de análisis de RFPs usando Gemini via Google AI API.
Trabaja con archivos locales extrayendo texto primero.
Soporta grounding para obtener tarifas de mercado actuales.
"""
import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Literal

from pypdf import PdfReader
from docx import Document
import io

from core.gcp.gemini_client import get_gemini_client
from core.services.parsers import get_parser_service
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from models.certification import Certification

logger = logging.getLogger(__name__)

# Cargar prompts
PROMPTS_DIR = Path(__file__).parent.parent.parent / "prompts"


def load_prompt(name: str) -> str:
    """Carga un prompt desde archivo."""
    prompt_path = PROMPTS_DIR / f"{name}.txt"
    if prompt_path.exists():
        return prompt_path.read_text(encoding="utf-8")
    
    # Prompt por defecto si no existe el archivo
    if name == "rfp_analysis":
        return DEFAULT_ANALYSIS_PROMPT
    elif name == "question_generation":
        return DEFAULT_QUESTIONS_PROMPT
    
    raise FileNotFoundError(f"Prompt not found: {prompt_path}")


DEFAULT_ANALYSIS_PROMPT = """
Eres un experto analista de RFPs (Request for Proposals) para TIVIT, una empresa líder en tecnología y servicios digitales en Latinoamérica.

Analiza el siguiente documento RFP y extrae la información estructurada en formato JSON con el siguiente schema:

{
    "title": "Título oficial del proyecto o licitación",
    "client_name": "Nombre COMPLETO del cliente/empresa",
    "client_acronym": "Siglas o abreviatura del cliente (ej: BCP, INDAP, BBVA)",
    "country": "País del cliente",
    "category": "Categoría del proyecto (infraestructura, desarrollo, cloud, seguridad, etc.)",
    "summary": "Resumen ejecutivo del RFP en máximo 3 párrafos",
    "scope": {
        "description": "Descripción detallada del alcance",
        "deliverables": ["Lista de entregables esperados"],
        "exclusions": ["Lo que NO está incluido"]
    },
    "requirements": {
        "technical": ["Requisitos técnicos"],
        "functional": ["Requisitos funcionales"],
        "certifications": ["Certificaciones requeridas"]
    },
    "budget": {
        "amount_min": null,
        "amount_max": null,
        "currency": "USD",
        "notes": "Notas sobre presupuesto"
    },
    "questions_deadline": "YYYY-MM-DD",
    "project_duration": "Duración estimada [Fuente: doc_X, Pag Y]",
    "evaluation_criteria": [
        {"criterion": "nombre", "weight": 0.0}
    ],
    "penalties": [
        {"description": "Descripción de multa [Fuente: doc_X, Pag Y]", "amount": "Monto", "is_high": true}
    ],
    "sla": [
        {"description": "Descripción del SLA [Fuente: doc_X, Pag Y]", "metric": "Metrica", "is_aggressive": true}
    ],
    "risks": [
        {"risk": "descripción [Fuente: doc_X, Pag Y]", "severity": "high/medium/low", "mitigation": "posible mitigación", "category": "financial/technical/legal"}
    ],
    "opportunities": ["Oportunidad [Fuente: doc_X, Pag Y]"],
    "recommendation_reasons": ["Razón 1 [Fuente: doc_X, Pag Y]", "Razón 2 [Fuente: doc_X, Pag Y]"],
    "confidence_score": 0.0,
    "recommendation": "GO o NO_GO con justificación breve"
}

Sé preciso y extrae toda la información relevante. Si algún campo no está disponible, usa null.
"""


DEFAULT_QUESTIONS_PROMPT = """
Eres un experto en desarrollo de negocios de TIVIT que prepara preguntas estratégicas para enviar al cliente antes de presentar una propuesta.

Basándote en el análisis del RFP, genera entre 10 y 15 preguntas inteligentes que:
1. Clarifiquen ambigüedades en los requisitos
2. Identifiquen oportunidades de valor agregado
3. Ayuden a preparar una propuesta más competitiva
4. Demuestren expertise y profesionalismo

Responde con un JSON array de objetos:
[
    {
        "question": "La pregunta completa",
        "category": "técnica|comercial|alcance|timeline|legal",
        "priority": "alta|media|baja",
        "context": "Por qué es importante esta pregunta",
        "why_important": "Cómo impacta en la propuesta"
    }
]

Ordena las preguntas por prioridad (alta primero).
"""


class RFPAnalyzerService:
    """Servicio para analizar RFPs con Gemini."""
    
    def __init__(self):
        self._gemini = None
        try:
            self.analysis_prompt = load_prompt("rfp_analysis")
        except FileNotFoundError:
            self.analysis_prompt = DEFAULT_ANALYSIS_PROMPT
        
        try:
            self.questions_prompt = load_prompt("question_generation")
        except FileNotFoundError:
            self.questions_prompt = DEFAULT_QUESTIONS_PROMPT
        
        try:
            self.certification_prompt = load_prompt("certification_analysis")
        except FileNotFoundError:
            self.certification_prompt = "Analiza este certificado y extrae: name, issuer, description, issue_date, expiry_date, scope en JSON."

        try:
            self.chapter_prompt = load_prompt("chapter_analysis")
        except FileNotFoundError:
            logger.warning("Chapter prompt not found, using fallback")
            self.chapter_prompt = "Analiza este capítulo de propuesta técnica y extrae: name (título del capítulo), description (breve resumen del contenido) en JSON."

        try:
            self.experience_prompt = load_prompt("experience_relevance")
        except FileNotFoundError:
            logger.warning("Experience prompt not found, using fallback")
            self.experience_prompt = """
            Actúa como un experto en preventa servicios TI.
            RESUMEN RFP: {rfp_summary}
            EXPERIENCIAS: {experiences_text}
            TAREA: Identifica relevantes (score > 0.4). JSON array output.
            """

    async def analyze_certification_content(self, content: bytes, filename: str) -> dict[str, Any]:
        """
        Analiza un documento de certificación.
        """
        logger.info(f"Analyzing certification: {filename}")
        document_text = self.extract_text(content, filename)
        
        if not document_text.strip():
            return {"name": filename, "description": "No text extracted"}
            
        result = await self.gemini.analyze_document(
            document_content=document_text,
            prompt=self.certification_prompt,
            analysis_mode="fast"
        )
        return result

    async def analyze_chapter_content(self, content: bytes, filename: str) -> dict[str, Any]:
        """
        Analiza un documento de capítulo.
        """
        logger.info(f"Analyzing chapter: {filename}")
        document_text = self.extract_text(content, filename)
        
        if not document_text.strip():
            return {"name": filename, "description": "No text extracted"}
            
        result = await self.gemini.analyze_document(
            document_content=document_text,
            prompt=self.chapter_prompt,
            analysis_mode="fast"
        )
        return result
    
    @property
    def gemini(self):
        """Lazy loading del cliente Gemini."""
        if self._gemini is None:
            self._gemini = get_gemini_client()
        return self._gemini
    
    def extract_text_from_pdf(self, content: bytes) -> str:
        """Extrae texto de un PDF."""
        try:
            reader = PdfReader(io.BytesIO(content))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise
    
    def extract_text_from_docx(self, content: bytes) -> str:
        """Extrae texto de un DOCX."""
        try:
            doc = Document(io.BytesIO(content))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # También extraer de tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise
    
    def extract_text(self, content: bytes, filename: str) -> str:
        """Extrae texto de un archivo según su extensión."""
        filename_lower = filename.lower()
        
        if filename_lower.endswith(".pdf"):
            return self.extract_text_from_pdf(content)
        elif filename_lower.endswith(".docx"):
            return self.extract_text_from_docx(content)
        else:
            # Asumir texto plano
            return content.decode("utf-8", errors="ignore")
    
    async def analyze_rfp_project(
        self,
        files: list[dict[str, Any]],  # List of {content: bytes, filename: str, type: str}
        analysis_mode: str = "balanced",
        db: AsyncSession | None = None,
    ) -> dict[str, Any]:
        """
        Analiza un PROYECTO RFP completo (múltiples archivos).
        Usa ingesta estructurada (Premium) y contexto XML.
        """
        logger.info(f"Starting Multi-File RFP Analysis. Files: {len(files)}")
        
        # 1. Construir Contexto Multimodal/Estructurado
        context_parts = []
        context_parts.append("<rfp_project_files>")
        
        parser = get_parser_service()
        
        for idx, file in enumerate(files, 1):
            filename = file["filename"]
            content = file["content"]
            file_type = file.get("type", "unknown")
            file_id = f"doc_{idx}"  # ID estable para citas
            
            logger.info(f"Processing File {idx}: {filename} ({file_type})")
            
            processed_content = ""
            
            if filename.lower().endswith(".xlsx") or filename.lower().endswith(".xls"):
                # Parseo estructurado de Excel
                processed_content = parser.parse_excel(content, filename)
                
            elif filename.lower().endswith(".pdf"):
                # Para PDFs, usamos extracción de texto simple por ahora para el contexto global
                # Opcional: Podríamos usar Gemini Vision separadamente si se necesitara
                processed_content = self.extract_text_from_pdf(content)
                
            elif filename.lower().endswith(".docx"):
                processed_content = self.extract_text_from_docx(content)
                
            else:
                try:
                    processed_content = content.decode("utf-8", errors="ignore")
                except:
                    processed_content = "[Contenido binario no legible]"
            
            # Envolver en XML tags con ID para citas
            context_parts.append(f"""
    <document id="{file_id}" name="{filename}" type="{file_type}">
    {processed_content}
    </document>""")

        context_parts.append("</rfp_project_files>")
        
        full_context = "\n".join(context_parts)
        
        # 2. Preparar Prompt Premium
        # Injectamos el contexto y las reglas de citación
        # The analysis_prompt is expected to already contain the premium instructions.
        # The full_context is passed as document_content.
        
        # 3. Llamada a Gemini (1.5 Pro maneja el contexto largo)
        result = await self.gemini.analyze_document(
            document_content=full_context,
            prompt=self.analysis_prompt,
            analysis_mode=analysis_mode,
        )
        
        if isinstance(result, list) and result:
             result = result[0]
             
        logger.info("Multi-File Analysis Completed")
        return result

    async def analyze_rfp_from_content(
        self, 
        content: bytes, 
        filename: str,
        analysis_mode: Literal["fast", "balanced", "deep"] = "balanced",
        use_grounding: bool = False,  # Disabled by default - produces unrealistic salaries
        db: AsyncSession | None = None,
    ) -> dict[str, Any]:
        """
        Analiza un RFP desde su contenido en bytes.
        
        ESTRATEGIA:
        - PDFs: Siempre usa método binario (mejor para tablas visuales, OCR automático)
        - DOCX: Usa extracción de texto
        - Grounding: Deshabilitado por defecto (produce tarifas irreales de USA/Europa)
        
        Args:
            content: Contenido del archivo en bytes
            filename: Nombre del archivo (para determinar tipo)
            analysis_mode: Modo de análisis (fast/balanced/deep)
            use_grounding: DEPRECATED - Producía tarifas 3x más altas que realidad LATAM
            db: Sesión de base de datos para obtener certificaciones
            
        Returns:
            Datos extraídos del RFP incluyendo team_estimation y cost_estimation
        """
        logger.info(f"Starting RFP analysis for: {filename}")
        logger.info(f"Analysis mode: {analysis_mode}")
        
        # Detectar tipo de archivo y usar estrategia apropiada
        filename_lower = filename.lower()
        
        # PDFs SIEMPRE usan método binario (mejor análisis de tablas)
        if filename_lower.endswith('.pdf'):
            logger.info("PDF detected - using BINARY method for better table recognition and OCR")
            return await self._analyze_pdf_binary(
                content, filename, analysis_mode, db
            )
        
        # DOCX y otros usan extracción de texto
        else:
            logger.info(f"{filename.split('.')[-1].upper()} detected - using TEXT extraction method")
            return await self._analyze_with_text_extraction(
                content, filename, analysis_mode, False, db  # Grounding always False
            )
    
    async def _analyze_pdf_binary(
        self,
        pdf_bytes: bytes,
        filename: str,
        analysis_mode: str,
        db: AsyncSession | None,
    ) -> dict[str, Any]:
        """
        Analiza PDF usando método binario (OCR de Gemini).
        Mantiene tablas, gráficos y layout visual intacto.
        """
        logger.info("Using PDF binary analysis method")
        
        # Preparar prompt con certificaciones
        prompt_to_use = self.analysis_prompt
        
        if db:
            try:
                result = await db.execute(
                    select(Certification).where(Certification.is_active == True)
                )
                certs = result.scalars().all()
                if certs:
                    cert_list_str = "\n".join([
                        f"- {c.name} (ID: {c.id}): {c.description[:100]}..."
                        for c in certs
                    ])
                    prompt_to_use = prompt_to_use.replace(
                        "{{available_certifications}}", cert_list_str
                    )
                    logger.info(f"Injected {len(certs)} certifications into prompt")
                else:
                    prompt_to_use = prompt_to_use.replace(
                        "{{available_certifications}}", 
                        "No hay certificaciones disponibles."
                    )
            except Exception as e:
                logger.error(f"Error fetching certifications: {e}")
                prompt_to_use = prompt_to_use.replace(
                    "{{available_certifications}}", 
                    "Error al recuperar certificaciones."
                )
        else:
            prompt_to_use = prompt_to_use.replace(
                "{{available_certifications}}", 
                "No disponible (sin conexión a DB)."
            )
        
        # Obtener configuración del modo
        from core.gcp.gemini_client import ANALYSIS_MODES
        mode_config = ANALYSIS_MODES.get(analysis_mode, ANALYSIS_MODES["balanced"])
        
        # Llamar al método de Gemini para PDFs binarios
        result = await self.gemini.analyze_pdf_bytes(
            pdf_bytes=pdf_bytes,
            prompt=prompt_to_use,
            temperature=mode_config["temperature"],
            max_output_tokens=mode_config["max_output_tokens"],
        )
        
        # Validar formato de respuesta
        if isinstance(result, list):
            if len(result) > 0 and isinstance(result[0], dict):
                logger.warning("Gemini returned a list instead of a dict. Using first item.")
                result = result[0]
            else:
                logger.error(f"Gemini returned unexpected list format: {result}")
                result = {"error": "Invalid response format from AI", "raw": result}
        
        logger.info(f"Binary PDF analysis completed: {result.get('title', 'Unknown')}")
        return result
    
    async def _analyze_with_text_extraction(
        self,
        content: bytes,
        filename: str,
        analysis_mode: str,
        use_grounding: bool,
        db: AsyncSession | None,
    ) -> dict[str, Any]:
        """
        Análisis con extracción de texto (para DOCX u otros formatos no-PDF).
        """
        logger.info("Using text extraction method")
        
        # Extraer texto del documento
        document_text = self.extract_text(content, filename)
        
        if not document_text.strip():
            logger.error("No text extracted from document")
            return {"error": "No se pudo extraer texto del documento"}
        
        logger.info(f"Extracted {len(document_text)} characters from document")
        
        # Preparar prompt con certificaciones
        prompt_to_use = self.analysis_prompt
        
        if db:
            try:
                result = await db.execute(select(Certification).where(Certification.is_active == True))
                certs = result.scalars().all()
                if certs:
                    cert_list_str = "\n".join([f"- {c.name} (ID: {c.id}): {c.description[:100]}..." for c in certs])
                    prompt_to_use = prompt_to_use.replace("{{available_certifications}}", cert_list_str)
                    logger.info(f"Injected {len(certs)} certifications into prompt")
                else:
                    prompt_to_use = prompt_to_use.replace("{{available_certifications}}", "No hay certificaciones disponibles.")
            except Exception as e:
                logger.error(f"Error fetching certifications for prompt: {e}")
                prompt_to_use = prompt_to_use.replace("{{available_certifications}}", "Error al recuperar certificaciones.")
        else:
            prompt_to_use = prompt_to_use.replace("{{available_certifications}}", "No disponible (sin conexión a DB).")

        # Grounding disabled - produces unrealistic USA/Europe salaries
        if use_grounding:
            logger.warning(
                "Grounding was requested but is DISABLED (produces 2-3x higher salaries than LATAM reality). "
                "Using standard analysis with AI knowledge base."
            )
        
        # Siempre usar análisis estándar (sin grounding)
        result = await self.gemini.analyze_document(
            document_content=document_text,
            prompt=prompt_to_use,
            analysis_mode=analysis_mode,
        )
        
        if isinstance(result, list):
            if len(result) > 0 and isinstance(result[0], dict):
                 logger.warning("Gemini returned a list instead of a dict. Using first item.")
                 result = result[0]
            else:
                 logger.error(f"Gemini returned unexpected list format: {result}")
                 result = {"error": "Invalid response format from AI", "raw": result}
        
        logger.info(f"Text extraction analysis completed: {result.get('title', 'Unknown Title')}")
        return result
    
    async def analyze_rfp(self, gcs_uri: str, use_grounding: bool = True, db: AsyncSession | None = None) -> dict[str, Any]:
        """
        Analiza un RFP desde GCS o local.
        
        Args:
            gcs_uri: URI del archivo en Cloud Storage o local://
            use_grounding: Si True, usa Google Search para tarifas de mercado
            db: Sesión de DB opcional
            
        Returns:
            Datos extraídos del RFP
        """
        logger.info(f"Starting RFP analysis: {gcs_uri}")
        
        # Si es un URI local, leer el archivo
        if gcs_uri.startswith("local://"):
            from core.storage import get_storage_service
            storage = get_storage_service()
            content = storage.download_file(gcs_uri)
            filename = gcs_uri.split("/")[-1]
            return await self.analyze_rfp_from_content(
                content, 
                filename, 
                use_grounding=use_grounding,
                db=db
            )
        
        # Para GCS, descargar y analizar
        from core.storage import get_storage_service
        storage = get_storage_service()
        content = storage.download_file(gcs_uri)
        filename = gcs_uri.split("/")[-1]
        return await self.analyze_rfp_from_content(
            content, 
            filename, 
            use_grounding=use_grounding,
            db=db
        )
    
    async def generate_questions(self, rfp_data: dict[str, Any]) -> list[dict[str, Any]]:
        """
        Genera preguntas basadas en el análisis del RFP.
        
        Args:
            rfp_data: Datos extraídos del análisis
            
        Returns:
            Lista de preguntas generadas
        """
        logger.info("Generating questions for RFP")
        
        questions = await self.gemini.generate_questions(
            rfp_data=rfp_data,
            prompt=self.questions_prompt,
            temperature=0.3,
        )
        
        logger.info(f"Generated {len(questions)} questions")
        return questions
    
    def extract_indexed_fields(self, extracted_data: dict[str, Any]) -> dict[str, Any]:
        """
        Extrae campos para indexación en la base de datos.
        
        Args:
            extracted_data: Datos completos extraídos
            
        Returns:
            Campos para actualizar en el modelo
        """
        # Manejar el campo budget de manera robusta
        budget_data = extracted_data.get("budget")
        if isinstance(budget_data, dict):
            budget = budget_data
        elif isinstance(budget_data, list) and len(budget_data) > 0 and isinstance(budget_data[0], dict):
            # Si es una lista, tomar el primer elemento
            budget = budget_data[0]
        else:
            # Por defecto, diccionario vacío
            budget = {}
        
        # Parsear fechas
        proposal_deadline = None
        questions_deadline = None
        
        if extracted_data.get("proposal_deadline"):
            try:
                proposal_deadline = datetime.strptime(
                    extracted_data["proposal_deadline"], "%Y-%m-%d"
                ).date()
            except ValueError:
                pass
        
        if extracted_data.get("questions_deadline"):
            try:
                questions_deadline = datetime.strptime(
                    extracted_data["questions_deadline"], "%Y-%m-%d"
                ).date()
            except ValueError:
                pass
        
        return {
            "title": extracted_data.get("title"),
            "client_name": extracted_data.get("client_name"),
            "client_acronym": extracted_data.get("client_acronym"),
            "country": extracted_data.get("country"),
            "category": extracted_data.get("category"),
            "summary": extracted_data.get("summary"),
            "budget_min": budget.get("amount_min") if isinstance(budget, dict) else None,
            "budget_max": budget.get("amount_max") if isinstance(budget, dict) else None,
            "currency": budget.get("currency", "USD") if isinstance(budget, dict) else "USD",
            "proposal_deadline": proposal_deadline,
            "questions_deadline": questions_deadline,
            "project_duration": extracted_data.get("project_duration"),
            "confidence_score": extracted_data.get("confidence_score"),
            "recommendation": extracted_data.get("recommendation"),
        }


    async def analyze_experience_relevance(
        self, 
        rfp_summary: str,
        experiences: list[dict[str, Any]]
    ) -> list[dict[str, Any]]:
        """
        Analiza la relevancia de las experiencias para un RFP.
        Ref: Step Id: 2289
        """
        if not experiences:
            return []

        # Prepare prompt
        experiences_text = "\n".join([
            f"ID: {exp['id']}\nCliente: {exp.get('propietario_servicio')}\nDescripción: {exp.get('descripcion_servicio')}\nMonto: {exp.get('monto_final')}\n---"
            for exp in experiences
        ])

        # Prepare prompt using the loaded template
        try:
            prompt = self.experience_prompt.format(
                rfp_summary=rfp_summary,
                experiences_text=experiences_text
            )
        except Exception as e:
            logger.error(f"Error formatting prompt: {e}")
            # Fallback simple prompt
            prompt = f"Analiza relevancia RFP: {rfp_summary} vs Experiencias: {experiences_text}. JSON array."

        try:
            # Use the new helper method directly
            ai_recommendations = await self.gemini.generate_json(prompt)
            
            # Ensure it's a list
            if isinstance(ai_recommendations, dict):
                ai_recommendations = [ai_recommendations]
            elif not isinstance(ai_recommendations, list):
                logger.warning(f"AI returned unexpected format: {type(ai_recommendations)}")
                ai_recommendations = []
            
            # --- POST-PROCESSING: Fill in missing items ---
            # The AI might filter out items. We need to ensure ALL experiences have a score for the frontend.
            
            # 1. Create a map of updated scores from AI
            ai_map = {rec.get("experience_id"): rec for rec in ai_recommendations if rec.get("experience_id")}
            
            final_recommendations = []
            
            # 2. Iterate over ALL original experiences
            for exp in experiences:
                exp_id = str(exp["id"])
                
                if exp_id in ai_map:
                    # Use AI's evaluation
                    final_recommendations.append(ai_map[exp_id])
                else:
                    # Provide default low score for items the AI ignored
                    final_recommendations.append({
                        "experience_id": exp_id,
                        "score": 0.05,
                        "reason": "No seleccionada por IA como relevante para este RFP."
                    })
            
            logger.info(f"AI returned {len(ai_recommendations)} relevant items. Augmented to {len(final_recommendations)} total items.")
            return final_recommendations

        except Exception as e:
            logger.error(f"Error in analyze_experience_relevance: {e}")
            logger.error(f"Stack trace:", exc_info=True)
            return []

    async def analyze_chapter_relevance(
        self, 
        rfp_summary: str,
        chapters: list[dict[str, Any]]
    ) -> list[dict[str, Any]]:
        """
        Analiza la relevancia de los capítulos para un RFP.
        """
        if not chapters:
            return []

        # Prepare prompt
        chapters_text = "\n".join([
            f"ID: {chap['id']}\nNombre: {chap['name']}\nDescripción: {chap.get('description', '')}\n---"
            for chap in chapters
        ])

        prompt = f"""
        Actúa como un experto en redacción de propuestas técnicas.
        
        RESUMEN DEL RFP:
        {rfp_summary}

        CAPÍTULOS DISPONIBLES (Biblioteca de Contenidos):
        {chapters_text}

        TAREA:
        Evalúa qué capítulos son útiles para responder a este RFP específico.
        Genera un JSON list con objetos: {{ "chapter_id": "UUID", "score": 0.0-1.0, "reason": "Breve justificación (max 10 palabras)" }}
        
        CRITERIOS:
        - Score > 0.8: Capítulo Esencial (ej: Metodología, Equipo, si el RFP lo pide explícitamente).
        - Score > 0.5: Capítulo Útil (Complementario).
        - Score < 0.5: Irrelevante.
        - Se estricto. Usa SOLO los IDs provistos.
        """

        try:
            recommendations = await self.gemini.generate_json(prompt)
            
            if isinstance(recommendations, dict):
                recommendations = [recommendations]
            elif not isinstance(recommendations, list):
                return []
            
            return recommendations

        except Exception as e:
            logger.error(f"Error in analyze_chapter_relevance: {e}")
            return []
            
# Singleton instance
_analyzer_service: RFPAnalyzerService | None = None


def get_analyzer_service() -> RFPAnalyzerService:
    """Get or create analyzer service singleton."""
    global _analyzer_service
    if _analyzer_service is None:
        _analyzer_service = RFPAnalyzerService()
    return _analyzer_service
