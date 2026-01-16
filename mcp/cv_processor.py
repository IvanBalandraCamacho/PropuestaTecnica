"""
CV Processor - Extrae texto de CVs (PDF/DOCX) y los prepara para indexacion.

Proceso:
1. Lee el archivo (PDF o DOCX)
2. Extrae el texto por pagina
3. [NUEVO] Extrae texto de imagenes usando Gemini Vision (certificaciones, badges, etc.)
4. Divide el texto en chunks con overlap para busqueda semantica
5. Retorna lista de chunks listos para vectorizar

Version 2.0 - Ahora con soporte Vision OCR para imagenes en CVs
"""

import re
from pathlib import Path
from typing import List, Optional, Tuple, Dict
from dataclasses import dataclass
import logging

logger = logging.getLogger(__name__)

# Imports opcionales para procesamiento de documentos
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logger.warning("PyMuPDF no instalado. Instalar con: pip install PyMuPDF")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx no instalado. Instalar con: pip install python-docx")

# Vision OCR con Gemini (opcional)
try:
    from cv_vision import process_cv_with_vision
    VISION_AVAILABLE = True
except ImportError:
    VISION_AVAILABLE = False
    logger.info("Vision OCR no disponible (cv_vision.py no encontrado)")


@dataclass
class CVChunk:
    """
    Fragmento de un CV para indexacion.
    
    Attributes:
        matricula: Matricula del colaborador (vinculada via cv_matcher)
        chunk_id: ID secuencial del chunk dentro del CV
        text: Texto del fragmento
        page_num: Numero de pagina (None para DOCX)
        cv_filename: Nombre del archivo original
    """
    matricula: str
    chunk_id: int
    text: str
    page_num: Optional[int]
    cv_filename: str


class CVProcessor:
    """
    Procesa CVs y los prepara para busqueda semantica.
    
    Caracteristicas:
    - Soporta PDF y DOCX
    - Chunkeriza con overlap para mejor contexto
    - Limpia y normaliza el texto
    - [NUEVO] Extrae texto de imagenes con Gemini Vision (certificaciones, badges, etc.)
    """
    
    def __init__(
        self, 
        cvs_folder: Path, 
        chunk_size: int = 500, 
        overlap: int = 100,
        use_vision: bool = True  # Nuevo parametro
    ):
        """
        Inicializa el procesador.
        
        Args:
            cvs_folder: Carpeta donde estan los CVs
            chunk_size: Tamano maximo de cada chunk en caracteres
            overlap: Caracteres de overlap entre chunks consecutivos
            use_vision: Si usar Gemini Vision para extraer texto de imagenes
        """
        self.cvs_folder = cvs_folder
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.use_vision = use_vision and VISION_AVAILABLE
    
    def extract_text_from_pdf(self, filepath: Path) -> List[Tuple[int, str]]:
        """
        Extrae texto de un PDF.
        
        Args:
            filepath: Ruta al archivo PDF
            
        Returns:
            Lista de tuplas (numero_pagina, texto)
        """
        if not PYMUPDF_AVAILABLE:
            logger.error("PyMuPDF no disponible para procesar PDFs")
            return []
        
        pages = []
        try:
            doc = fitz.open(filepath)
            for page_num, page in enumerate(doc, 1):
                text = page.get_text("text")
                if text and text.strip():
                    pages.append((page_num, text))
            doc.close()
            logger.debug(f"PDF procesado: {filepath.name} ({len(pages)} paginas)")
        except Exception as e:
            logger.error(f"Error procesando PDF {filepath}: {e}")
        
        return pages
    
    def extract_text_from_docx(self, filepath: Path) -> List[Tuple[int, str]]:
        """
        Extrae texto de un DOCX.
        
        Nota: DOCX no tiene concepto de "paginas" como PDF,
        asi que se retorna todo el contenido como pagina 1.
        
        Args:
            filepath: Ruta al archivo DOCX
            
        Returns:
            Lista con una tupla (1, texto_completo)
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx no disponible para procesar DOCX")
            return []
        
        try:
            doc = Document(filepath)
            paragraphs = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Tambien extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)
            
            full_text = "\n".join(paragraphs)
            
            if full_text:
                logger.debug(f"DOCX procesado: {filepath.name} ({len(full_text)} chars)")
                return [(1, full_text)]
            
        except Exception as e:
            logger.error(f"Error procesando DOCX {filepath}: {e}")
        
        return []
    
    def clean_text(self, text: str) -> str:
        """
        Limpia y normaliza el texto.
        
        - Remueve caracteres de control
        - Normaliza espacios y saltos de linea
        - Remueve lineas muy cortas (probablemente basura)
        """
        if not text:
            return ""
        
        # Remover caracteres de control excepto newlines
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        
        # Normalizar saltos de linea
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remover lineas que son solo espacios o muy cortas
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            # Mantener lineas con al menos 3 caracteres alfabeticos
            if line and len(re.sub(r'[^a-zA-Z]', '', line)) >= 3:
                cleaned_lines.append(line)
        
        text = '\n'.join(cleaned_lines)
        
        # Normalizar espacios multiples
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Normalizar multiples newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    def chunk_text(self, text: str) -> List[str]:
        """
        Divide el texto en chunks con overlap.
        
        Intenta cortar en limites naturales (punto, newline, espacio).
        
        Args:
            text: Texto a dividir
            
        Returns:
            Lista de chunks
        """
        # Limpiar y normalizar espacios para chunking
        text = re.sub(r'\s+', ' ', text).strip()
        
        if not text:
            return []
        
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Si no es el ultimo chunk, buscar punto de corte natural
            if end < len(text):
                # Buscar en orden de preferencia: punto+espacio, newline, espacio
                best_cut = None
                for separator in ['. ', '\n', ', ', ' ']:
                    cut_pos = text.rfind(separator, start + self.chunk_size // 2, end)
                    if cut_pos > start:
                        best_cut = cut_pos + len(separator)
                        break
                
                if best_cut:
                    end = best_cut
            
            chunk = text[start:end].strip()
            
            if chunk:
                chunks.append(chunk)
            
            # Siguiente chunk empieza con overlap
            start = end - self.overlap
            
            # Evitar loop infinito
            if start >= len(text) - 10:
                break
        
        return chunks
    
    def process_cv(self, filepath: Path, matricula: str) -> List[CVChunk]:
        """
        Procesa un CV completo y retorna sus chunks.
        
        Args:
            filepath: Ruta al archivo CV
            matricula: Matricula del colaborador
            
        Returns:
            Lista de CVChunk listos para indexar
        """
        ext = filepath.suffix.lower()
        
        # Extraer texto segun formato
        if ext == '.pdf':
            pages = self.extract_text_from_pdf(filepath)
        elif ext in ['.docx', '.doc']:
            pages = self.extract_text_from_docx(filepath)
        else:
            logger.warning(f"Formato no soportado: {filepath}")
            return []
        
        if not pages:
            logger.warning(f"No se extrajo texto de: {filepath}")
            return []
        
        chunks = []
        chunk_id = 0
        
        # Procesar texto normal de cada pagina
        for page_num, page_text in pages:
            # Limpiar texto
            cleaned_text = self.clean_text(page_text)
            
            if not cleaned_text:
                continue
            
            # Dividir en chunks
            for chunk_text in self.chunk_text(cleaned_text):
                if len(chunk_text) < 20:  # Ignorar chunks muy pequenos
                    continue
                    
                chunks.append(CVChunk(
                    matricula=matricula,
                    chunk_id=chunk_id,
                    text=chunk_text,
                    page_num=page_num,
                    cv_filename=filepath.name
                ))
                chunk_id += 1
        
        # [NUEVO] Extraer texto de imagenes con Gemini Vision
        if self.use_vision:
            try:
                vision_text = process_cv_with_vision(filepath)
                if vision_text:
                    logger.info(f"  Vision OCR: {len(vision_text)} chars extraidos de imagenes")
                    
                    # Limpiar y chunkerizar el texto de imagenes
                    cleaned_vision = self.clean_text(vision_text)
                    if cleaned_vision:
                        for chunk_text in self.chunk_text(cleaned_vision):
                            if len(chunk_text) < 20:
                                continue
                            
                            chunks.append(CVChunk(
                                matricula=matricula,
                                chunk_id=chunk_id,
                                text=f"[IMAGEN] {chunk_text}",  # Prefijo para identificar
                                page_num=0,  # Imagenes no tienen pagina especifica en este contexto
                                cv_filename=filepath.name
                            ))
                            chunk_id += 1
            except Exception as e:
                logger.warning(f"Error en Vision OCR para {filepath.name}: {e}")
        
        return chunks
    
    def process_all(self, mapping: Dict[str, str]) -> List[CVChunk]:
        """
        Procesa todos los CVs usando el mapping nombre -> matricula.
        
        Args:
            mapping: Dict {filename: matricula}
            
        Returns:
            Lista de todos los CVChunk
        """
        all_chunks = []
        processed = 0
        skipped = 0
        
        if not self.cvs_folder.exists():
            logger.warning(f"Carpeta no existe: {self.cvs_folder}")
            return []
        
        cv_files = [
            f for f in self.cvs_folder.iterdir()
            if f.suffix.lower() in ['.pdf', '.docx', '.doc']
        ]
        
        logger.info(f"Procesando {len(cv_files)} CVs...")
        
        for filepath in cv_files:
            matricula = mapping.get(filepath.name)
            
            if not matricula:
                logger.debug(f"Sin matricula para: {filepath.name}")
                skipped += 1
                continue
            
            chunks = self.process_cv(filepath, matricula)
            
            if chunks:
                all_chunks.extend(chunks)
                processed += 1
                logger.info(f"  {filepath.name[:40]:<40} -> {len(chunks)} chunks")
            else:
                logger.warning(f"  {filepath.name[:40]:<40} -> Sin contenido")
        
        logger.info(f"Procesados: {processed}, Omitidos: {skipped}, Total chunks: {len(all_chunks)}")
        
        return all_chunks
    
    def get_cv_path(self, filename: str) -> Optional[Path]:
        """
        Obtiene la ruta completa de un CV por su nombre.
        
        Args:
            filename: Nombre del archivo
            
        Returns:
            Path completo o None si no existe
        """
        cv_path = self.cvs_folder / filename
        return cv_path if cv_path.exists() else None


def check_dependencies() -> Dict[str, bool]:
    """
    Verifica que las dependencias esten instaladas.
    
    Returns:
        Dict con estado de cada dependencia
    """
    return {
        "PyMuPDF": PYMUPDF_AVAILABLE,
        "python-docx": DOCX_AVAILABLE,
        "Vision OCR": VISION_AVAILABLE
    }
